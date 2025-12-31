"""Apply updates directly to DOCX files."""

import shutil
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

# XML namespaces
NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}


def iter_all_paragraphs(doc: Document):
    """Iterate all paragraphs in document, including those in tables."""
    # Main body paragraphs
    for para in doc.paragraphs:
        yield para

    # Paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def find_paragraph_by_id(doc: Document, para_id: str) -> Optional[Paragraph]:
    """Find a paragraph by its w14:paraId."""
    for para in iter_all_paragraphs(doc):
        p_element = para._element
        element_id = p_element.get(qn("w14:paraId"))
        if element_id == para_id:
            return para
    return None


def find_paragraph_element_by_id(doc: Document, para_id: str):
    """Find paragraph XML element by ID."""
    for para in iter_all_paragraphs(doc):
        p_element = para._element
        element_id = p_element.get(qn("w14:paraId"))
        if element_id == para_id:
            return p_element
    return None


def extract_para_id_from_target(target_element: str) -> str:
    """
    Extract the paragraph ID from target_element string.
    Format: 'p-{paraId}' or 'li-{paraId}'
    """
    if target_element.startswith("p-"):
        return target_element[2:]
    elif target_element.startswith("li-"):
        return target_element[3:]
    else:
        raise ValueError(f"Cannot extract paragraph ID from: {target_element}")


def apply_replace(doc: Document, target_element: str, content: str) -> None:
    """Replace paragraph content."""
    para_id = extract_para_id_from_target(target_element)
    para = find_paragraph_by_id(doc, para_id)

    if para is None:
        raise ValueError(f"Paragraph not found: {target_element} (paraId: {para_id})")

    # Clear existing runs
    for run in para.runs:
        run.clear()

    # Clear all child elements except paragraph properties
    p_element = para._element
    for child in list(p_element):
        if not child.tag.endswith("}pPr"):
            p_element.remove(child)

    # Add new content as a single run
    para.add_run(content)


def apply_delete(doc: Document, target_element: str) -> None:
    """Delete a paragraph."""
    para_id = extract_para_id_from_target(target_element)
    p_element = find_paragraph_element_by_id(doc, para_id)

    if p_element is None:
        raise ValueError(f"Paragraph not found: {target_element} (paraId: {para_id})")

    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)


def create_paragraph_element(content: str, para_id: str) -> etree._Element:
    """Create a new paragraph XML element."""
    nsmap = {
        "w": NAMESPACES["w"],
        "w14": NAMESPACES["w14"],
    }

    p = etree.Element(qn("w:p"), nsmap=nsmap)
    p.set(qn("w14:paraId"), para_id)
    p.set(qn("w14:textId"), para_id[::-1])  # Use reversed ID for textId

    # Create run with text
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = content

    return p


def apply_insert_before(doc: Document, target_element: str, content: str) -> None:
    """Insert a new paragraph before target."""
    para_id = extract_para_id_from_target(target_element)
    p_element = find_paragraph_element_by_id(doc, para_id)

    if p_element is None:
        raise ValueError(f"Paragraph not found: {target_element} (paraId: {para_id})")

    # Generate new paragraph ID
    import secrets

    new_para_id = secrets.token_hex(4).upper()

    # Create new paragraph
    new_p = create_paragraph_element(content, new_para_id)

    # Insert before target
    p_element.addprevious(new_p)


def apply_insert_after(doc: Document, target_element: str, content: str) -> None:
    """Insert a new paragraph after target."""
    para_id = extract_para_id_from_target(target_element)
    p_element = find_paragraph_element_by_id(doc, para_id)

    if p_element is None:
        raise ValueError(f"Paragraph not found: {target_element} (paraId: {para_id})")

    # Generate new paragraph ID
    import secrets

    new_para_id = secrets.token_hex(4).upper()

    # Create new paragraph
    new_p = create_paragraph_element(content, new_para_id)

    # Insert after target
    p_element.addnext(new_p)


def apply_update_to_docx(doc: Document, update: dict) -> None:
    """Apply a single update to the DOCX document."""
    update_type = update["type"]
    target_element = update["target_element"]
    content = update.get("content", "")

    # Skip table-related updates for now (tables don't have simple paragraph IDs)
    if target_element.startswith(("tbl-", "tr-", "td-")):
        # TODO: Implement table updates
        return

    if update_type == "replace":
        apply_replace(doc, target_element, content)
    elif update_type == "insert_before":
        apply_insert_before(doc, target_element, content)
    elif update_type == "insert_after":
        apply_insert_after(doc, target_element, content)
    elif update_type == "delete":
        apply_delete(doc, target_element)
    else:
        raise ValueError(f"Unknown update type: {update_type}")


def apply_updates_to_docx(
    base_docx_path: Path,
    output_path: Path,
    updates: list[dict],
) -> None:
    """
    Apply all updates to a DOCX file.

    Args:
        base_docx_path: Path to the base DOCX (with IDs)
        output_path: Path for the output DOCX
        updates: List of update commands
    """
    # Copy base to output first
    shutil.copy2(base_docx_path, output_path)

    # Open the copied document
    doc = Document(str(output_path))

    # Apply each update
    for update in updates:
        try:
            apply_update_to_docx(doc, update)
        except ValueError as e:
            # Log warning but continue with other updates
            print(f"Warning: Skipping update - {e}")

    # Save the modified document
    doc.save(str(output_path))
